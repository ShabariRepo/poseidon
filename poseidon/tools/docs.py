"""Office documents — the files real teams actually work with.
read_document: xlsx / docx / pdf → text and tables.
edit_spreadsheet: set cells / append rows in xlsx (creates the file if new);
approval-gated and versioned like every other write.
"""
import asyncio

from .files import resolve_path

MAX_ROWS = 400
MAX_TEXT = 20_000


def _read_xlsx(path):
    from openpyxl import load_workbook
    wb = load_workbook(path, data_only=True, read_only=True)
    out = {}
    for ws in wb.worksheets[:6]:
        rows = []
        for row in ws.iter_rows(values_only=True):
            rows.append([("" if c is None else str(c))[:200] for c in row])
            if len(rows) >= MAX_ROWS:
                rows.append(["… truncated …"])
                break
        out[ws.title] = rows
    return {"type": "spreadsheet", "sheets": out}


def _read_docx(path):
    from docx import Document
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for tbl in doc.tables[:10]:
        for row in tbl.rows[:60]:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
    return {"type": "document", "text": "\n".join(parts)[:MAX_TEXT]}


def _read_pdf(path):
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    text = "\n".join((pg.extract_text() or "") for pg in reader.pages[:25])
    return {"type": "pdf", "pages": len(reader.pages), "text": text[:MAX_TEXT]}


def _read_document_sync(workdir, rel):
    path = resolve_path(workdir, rel)
    if not path.is_file():
        return {"error": f"not a file: {rel}"}
    suffix = path.suffix.lower()
    if suffix == ".xlsx":
        return _read_xlsx(path)
    if suffix == ".docx":
        return _read_docx(path)
    if suffix == ".pdf":
        return _read_pdf(path)
    return {"error": f"unsupported format {suffix} — use read_file for plain text"}


def _edit_spreadsheet_sync(workdir, rel, updates, append_rows, sheet):
    from openpyxl import Workbook, load_workbook
    path = resolve_path(workdir, rel)
    if path.suffix.lower() != ".xlsx":
        return {"error": "edit_spreadsheet only handles .xlsx"}
    if path.is_file():
        wb = load_workbook(path)
    else:
        wb = Workbook()
        if sheet:
            wb.active.title = sheet
    ws = wb[sheet] if sheet and sheet in wb.sheetnames else wb.active
    changed = 0
    for u in (updates or [])[:200]:
        if u.get("cell"):
            ws[u["cell"]] = u.get("value", "")
            changed += 1
    for row in (append_rows or [])[:200]:
        ws.append(row if isinstance(row, list) else [row])
        changed += 1
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)
    return {"ok": True, "changes": changed, "sheet": ws.title}


async def read_document(args, ctx):
    return await asyncio.to_thread(_read_document_sync, ctx["workdir"], args["path"])


async def edit_spreadsheet(args, ctx):
    return await asyncio.to_thread(
        _edit_spreadsheet_sync, ctx["workdir"], args["path"],
        args.get("updates"), args.get("append_rows"), args.get("sheet"))
